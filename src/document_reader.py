import os
import fitz
import docx
import pytesseract

from PIL import Image
from config import TESSERACT_PATH, PDF_TEXT_THRESHOLD


# ---------------------------------
# Tesseract Configuration
# ---------------------------------
# Use path from config if provided, otherwise
# let the system find tesseract via PATH

if TESSERACT_PATH:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH


# =============================================================
# PDF Extraction
# =============================================================

def _is_page_scanned(page):
    """
    Returns True if a PDF page has no meaningful text
    and should be treated as a scanned image page.
    """
    text = page.get_text().strip()
    return len(text) < PDF_TEXT_THRESHOLD


def _ocr_pdf_page(page):
    """
    Converts a single PDF page to an image and runs
    Tesseract OCR on it.

    Returns extracted text string.
    """

    # Render page to image at 300 DPI for good OCR accuracy
    matrix = fitz.Matrix(300 / 72, 300 / 72)
    pixmap = page.get_pixmap(matrix=matrix)

    # Convert pixmap to PIL Image without saving to disk
    image = Image.frombytes(
        "RGB",
        [pixmap.width, pixmap.height],
        pixmap.samples
    )

    return pytesseract.image_to_string(image)


def extract_pdf_text(file_path):
    """
    Extracts text from a PDF file.

    Handles three scenarios:
      1. Digital PDF          — all pages have text layer
      2. Scanned PDF          — no pages have text layer
      3. Mixed PDF            — some pages scanned, some digital

    For scanned pages, OCR fallback is applied automatically.
    The caller receives the full text regardless of PDF type.
    """

    full_text = []
    pdf = fitz.open(file_path)

    for page_number, page in enumerate(pdf, start=1):

        page_text = page.get_text().strip()

        if _is_page_scanned(page):

            # Scanned page — run OCR fallback
            print(
                f"  [OCR] Page {page_number} appears scanned "
                f"— applying OCR"
            )

            ocr_text = _ocr_pdf_page(page)
            full_text.append(ocr_text)

        else:

            # Digital page — use extracted text directly
            full_text.append(page_text)

    pdf.close()

    return "\n\n".join(full_text)


# =============================================================
# DOCX Extraction
# =============================================================

def _extract_table_text(table):
    """
    Extracts text from a DOCX table row by row.

    Each row is joined with tabs so the LLM can
    understand the column structure.
    Each row is on its own line.
    """

    rows = []

    for row in table.rows:

        cells = [cell.text.strip() for cell in row.cells]

        # Skip completely empty rows
        if any(cells):
            rows.append("\t".join(cells))

    return "\n".join(rows)


def extract_docx_text(file_path):
    """
    Extracts text from a DOCX file.

    Handles both paragraphs and tables in document order.
    Tables are critical in RFQ documents where product
    lists, quantities, and prices are structured as grids.
    """

    document = docx.Document(file_path)

    content = []

    # python-docx exposes document body elements in order.
    # Iterating document.element.body preserves reading order
    # across both paragraphs and tables.

    for element in document.element.body:

        # Paragraph element
        tag = element.tag.split("}")[-1]

        if tag == "p":

            # Find the matching paragraph object
            for paragraph in document.paragraphs:
                if paragraph._element is element:
                    text = paragraph.text.strip()
                    if text:
                        content.append(text)
                    break

        # Table element
        elif tag == "tbl":

            for table in document.tables:
                if table._element is element:
                    table_text = _extract_table_text(table)
                    if table_text:
                        content.append(table_text)
                    break

    return "\n".join(content)


# =============================================================
# Image Extraction (PNG, JPG, JPEG, TIFF, BMP)
# =============================================================

def extract_image_text(file_path):
    """
    Runs Tesseract OCR on an image file.

    Supported formats: PNG, JPG, JPEG, TIFF, TIF, BMP
    TIFF is common from corporate scanners and
    multifunction printers.
    """

    image = Image.open(file_path)

    return pytesseract.image_to_string(image)


# =============================================================
# Main Router
# =============================================================

def read_document(file_path):
    """
    Routes a file to the correct extraction function
    based on its extension.

    Returns extracted text as a string.
    Returns empty string for unsupported file types
    so the pipeline continues safely.

    Supported:
      .pdf            — digital and scanned, with OCR fallback
      .docx           — paragraphs and tables
      .png .jpg .jpeg — OCR via Tesseract
      .tiff .tif      — OCR via Tesseract (scanner output)
      .bmp            — OCR via Tesseract
    """

    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".pdf":
        return extract_pdf_text(file_path)

    elif extension == ".docx":
        return extract_docx_text(file_path)

    elif extension in [".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp"]:
        return extract_image_text(file_path)

    else:
        print(f"  [SKIP] Unsupported file type: {extension}")
        return ""